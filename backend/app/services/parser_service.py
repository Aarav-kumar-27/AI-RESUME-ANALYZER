import io
import re
import logging
import fitz  # PyMuPDF
from docx import Document
from typing import Optional

logger = logging.getLogger("ai_resume_analyzer.parser")

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB
MIN_EXTRACTED_TEXT_LENGTH = 20

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/x-pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/octet-stream"
}

# Custom Exceptions
class ParserError(Exception):
    """Base exception for parser errors."""
    pass

class UnsupportedFileTypeError(ParserError):
    """Raised when uploaded file format or content type is not supported."""
    pass

class FileTooLargeError(ParserError):
    """Raised when file size exceeds the 5 MB threshold."""
    pass

class EmptyDocumentError(ParserError):
    """Raised when document contains no readable text."""
    pass

class DocumentParseError(ParserError):
    """Raised when document is corrupted or unreadable."""
    pass


def clean_extracted_text(raw_text: str) -> str:
    """
    Clean extracted text preserving structure, bullet points, and headers.
    - Normalizes carriage returns and line endings to \\n
    - Collapses 3+ consecutive newlines to double newlines
    - Strips trailing/leading space per line while keeping line structure
    """
    if not raw_text:
        return ""

    # Normalize carriage returns and line endings
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")

    # Clean per-line trailing/leading whitespace while keeping non-empty line content
    lines = [line.strip() for line in text.split("\n")]
    
    # Rejoin lines
    text = "\n".join(lines)

    # Collapse 3 or more consecutive newlines into 2 (paragraph boundary)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Collapse multiple inline spaces/tabs to a single space (except newlines)
    text = re.sub(r"[ \t]+", " ", text)

    return text.strip()


def validate_file_header_magic_bytes(file_bytes: bytes, extension: str) -> bool:
    """Validate file content using magic byte signatures."""
    if extension == ".pdf":
        return file_bytes.startswith(b"%PDF-")
    elif extension == ".docx":
        # Zip file header signature PK\x03\x04
        return file_bytes.startswith(b"PK\x03\x04")
    return False


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes in-memory using PyMuPDF (fitz)."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if doc.is_encrypted:
            raise DocumentParseError("The PDF file is password protected or encrypted.")
        
        pages_text = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            if text:
                pages_text.append(text)
        
        doc.close()
        return "\n".join(pages_text)
    except DocumentParseError:
        raise
    except Exception as e:
        logger.error(f"PyMuPDF failed to parse PDF: {str(e)}")
        raise DocumentParseError("Failed to parse PDF document. The file may be corrupt or invalid.")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes in-memory using python-docx."""
    try:
        doc_stream = io.BytesIO(file_bytes)
        doc = Document(doc_stream)
        
        paragraphs_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                paragraphs_text.append(paragraph.text)
                
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    paragraphs_text.append(" | ".join(row_cells))
                    
        return "\n".join(paragraphs_text)
    except Exception as e:
        logger.error(f"python-docx failed to parse DOCX: {str(e)}")
        raise DocumentParseError("Failed to parse DOCX document. The file may be corrupt or invalid.")


def extract_resume_text(
    file_bytes: bytes,
    filename: str,
    content_type: Optional[str] = None
) -> str:
    """
    Main parser service interface for resume validation and text extraction.
    Accepts in-memory bytes, validates extension, MIME type, size, magic bytes, and extracts clean text.
    """
    # 1. Validate File Size
    if not file_bytes or len(file_bytes) == 0:
        raise EmptyDocumentError("Uploaded file is empty (0 bytes).")
        
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise FileTooLargeError(
            f"File size ({len(file_bytes) / (1024 * 1024):.2f} MB) exceeds the maximum allowed limit of 5 MB."
        )

    # 2. Validate Extension
    filename_lower = filename.lower() if filename else ""
    ext = None
    for allowed_ext in ALLOWED_EXTENSIONS:
        if filename_lower.endswith(allowed_ext):
            ext = allowed_ext
            break
            
    if not ext:
        raise UnsupportedFileTypeError(
            f"Unsupported file extension in '{filename}'. Only .pdf and .docx files are supported."
        )

    # 3. Validate MIME Type if provided
    if content_type and content_type.lower() not in ALLOWED_MIME_TYPES:
        logger.warning(f"Unexpected MIME type '{content_type}' for file '{filename}'")

    # 4. Validate Magic Bytes
    if not validate_file_header_magic_bytes(file_bytes, ext):
        raise UnsupportedFileTypeError(
            f"File content does not match expected signature for {ext.upper()} format."
        )

    # 5. Perform Extraction
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise UnsupportedFileTypeError(f"Unsupported format '{ext}'.")

    # 6. Clean and Validate Extracted Text
    cleaned_text = clean_extracted_text(raw_text)

    if len(cleaned_text) < MIN_EXTRACTED_TEXT_LENGTH:
        raise EmptyDocumentError(
            "Document contains no meaningful readable text (less than 20 characters extracted)."
        )

    return cleaned_text
