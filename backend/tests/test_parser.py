import io
import fitz
import pytest
from docx import Document

from app.services.parser_service import (
    extract_resume_text,
    clean_extracted_text,
    UnsupportedFileTypeError,
    FileTooLargeError,
    EmptyDocumentError,
    DocumentParseError
)

def create_sample_pdf_bytes(text: str = "Jane Doe\nSoftware Engineer with 5 years experience in Python and React.") -> bytes:
    """Generate in-memory PDF file bytes using PyMuPDF."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    pdf_bytes = doc.write()
    doc.close()
    return pdf_bytes

def create_sample_docx_bytes(text: str = "John Smith\nSenior Full Stack Developer specializing in FastAPI and React.") -> bytes:
    """Generate in-memory DOCX file bytes using python-docx."""
    doc = Document()
    doc.add_heading("John Smith", level=1)
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# 1. Valid Document Parsing Tests
def test_extract_resume_text_valid_pdf():
    pdf_bytes = create_sample_pdf_bytes("Alice Developer\nExperienced Python Engineer building microservices.")
    text = extract_resume_text(pdf_bytes, "resume.pdf", "application/pdf")
    assert "Alice Developer" in text
    assert "Python Engineer" in text

def test_extract_resume_text_valid_docx():
    docx_bytes = create_sample_docx_bytes("Bob Architect\nLead Cloud Engineer with AWS and Kubernetes skills.")
    text = extract_resume_text(docx_bytes, "resume.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "Bob Architect" in text
    assert "Cloud Engineer" in text

# 2. Unsupported File Types
def test_extract_resume_text_unsupported_txt():
    txt_bytes = b"Jane Doe\nPlain text resume content."
    with pytest.raises(UnsupportedFileTypeError) as exc_info:
        extract_resume_text(txt_bytes, "resume.txt", "text/plain")
    assert "Unsupported file extension" in str(exc_info.value)

def test_extract_resume_text_unsupported_image():
    png_bytes = b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
    with pytest.raises(UnsupportedFileTypeError) as exc_info:
        extract_resume_text(png_bytes, "resume.png", "image/png")
    assert "Unsupported file extension" in str(exc_info.value)

# 3. File Size Validation
def test_extract_resume_text_oversized_file():
    # 5 MB + 1 byte
    large_bytes = b"%PDF-1.4 " + (b"A" * (5 * 1024 * 1024 + 1))
    with pytest.raises(FileTooLargeError) as exc_info:
        extract_resume_text(large_bytes, "huge.pdf", "application/pdf")
    assert "exceeds the maximum allowed limit of 5 MB" in str(exc_info.value)

# 4. Empty Files
def test_extract_resume_text_empty_file():
    with pytest.raises(EmptyDocumentError) as exc_info:
        extract_resume_text(b"", "empty.pdf", "application/pdf")
    assert "Uploaded file is empty" in str(exc_info.value)

# 5. Corrupt / Unreadable Files
def test_extract_resume_text_corrupt_pdf():
    # File has .pdf extension and %PDF- header but contains invalid binary structure
    corrupt_pdf = b"%PDF-1.4 corrupt garbage data"
    with pytest.raises(DocumentParseError) as exc_info:
        extract_resume_text(corrupt_pdf, "corrupt.pdf", "application/pdf")
    assert "Failed to parse PDF document" in str(exc_info.value)

def test_extract_resume_text_corrupt_docx():
    # Zip magic bytes PK\x03\x04 but invalid zip content
    corrupt_docx = b"PK\x03\x04 corrupt zip structure"
    with pytest.raises(DocumentParseError) as exc_info:
        extract_resume_text(corrupt_docx, "corrupt.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "Failed to parse DOCX document" in str(exc_info.value)

# 6. Valid File with Insufficient / No Meaningful Text
def test_extract_resume_text_no_meaningful_text():
    # PDF with only a few characters (< 20 chars)
    short_pdf = create_sample_pdf_bytes("Hi")
    with pytest.raises(EmptyDocumentError) as exc_info:
        extract_resume_text(short_pdf, "short.pdf", "application/pdf")
    assert "no meaningful readable text" in str(exc_info.value)

# 7. Text Cleaning & Whitespace Normalization
def test_clean_extracted_text_normalization():
    raw_text = "  Jane Doe  \r\n\r\n\r\n  Software   Engineer  \n\n\n\n- Python\n- React  "
    cleaned = clean_extracted_text(raw_text)
    assert cleaned == "Jane Doe\n\nSoftware Engineer\n\n- Python\n- React"
